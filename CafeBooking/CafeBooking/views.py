from io import StringIO
from django.shortcuts import render, redirect
from .forms import *
from .decorators import admin_decorator, HttpResponseRedirect, authorized, head_decorator
from django.http import FileResponse, HttpResponse
import pandas as pd
from django.core.management import call_command
import json
from dateutil.relativedelta import relativedelta
import os
from PIL import Image, ImageDraw, ImageFont
import qrcode
import socket
from openpyxl.drawing.image import Image as ExcelImage
from openpyxl import Workbook
from django.db import connection
from django.db import transaction
from prometheus_client import CollectorRegistry, Counter, Gauge

def create_reservation(user: User, time: Time, status: Status, place: Type, date: datetime.date):
    """:user: object of class User \n
    :time: object of class Time \n
    :status: object of class Status \n
    :place: object of class Type \n
    :date: Object of class datetime.date \n
    Creates a reservation and saves it to database
    """
    
    reservation = Reservation()
    reservation.user = user
    reservation.time = time
    reservation.status = status
    reservation.type = place
    reservation.date = date
    reservation.save()


@admin_decorator
def delete_rs_archived(request):
    """Executes function called 'delete_archived' and redirects back to admin-page"""
    Reservation.objects.filter(is_archivated=True).delete()
    return redirect('/admin/')


@authorized
def book_qr_code(request, id):
    """
    :id: Id of a place\n
    Creates a 'Reservation' object and saves it to database\n
    After that redirects on error page or success page
    """
    place = Type.objects.get(pk=id)
    today = datetime.date.today()
    error = None
    time = "Утреннее" if datetime.datetime.now().time() < datetime.time(14, 0) \
        else "Вечернее" if datetime.datetime.now().time() > datetime.time(14, 0) \
            and datetime.datetime.now().time() < datetime.time(19, 0) else None
    
    if time is None:
        error = "На сегодня бронирование больше недоступно!"
        args = {'error': error}
        return render(request, "error.html", args)
    
    is_booked = Reservation.objects.exclude(is_archivated=True) \
                    .filter(time__time=time, date=today, type=place, status__status="Забронировано").exists()
    
    if is_booked:
        error = "Данное место уже забронировано!"
        args = {'error': error}
        return render(request, "error.html", args)
    
    id = request.session.get('id', -1)
    if id == -1:
        return redirect('/login/')
    user = User.objects.filter(id = id).first()
    
    create_reservation(user, Time.objects.filter(time=time).first(), 
                       Status.objects.filter(status="Забронировано").first(), place, today)
    
    return redirect('/success/')
    
    
def create_if_not_exists(dir_name: str):
    """
    :dir_name: Name of a directory to create\n
    Creates a directory if it doesn't exists in CWD
    """
    cwd = os.getcwd()
    if not os.path.exists(f'{cwd}\\{dir_name}'):
        path = os.path.join(cwd, dir_name)
        os.mkdir(path)


@head_decorator
def __download_qrcodes(request):
    """
    Creates an xlsx file with QR-code pictures in it.\n
    Returns file as a response.
    """
    id = request.session.get('id', -1)
    if id == -1:
        return redirect('/login/')
    user = User.objects.filter(id = id).first()
    cwd = os.getcwd()
    
    create_if_not_exists("qr-codes")
        
    offices = list(Office.objects.all())
    
    for office in offices:
        data = []
        
        create_if_not_exists(f"qr-codes\\office {office.id}")
            
        places = list(Type.objects.filter(place__office=office).all())
        for place in places:
            url = f"http://{socket.gethostbyname(socket.gethostname())}:8000/qr_book/{place.pk}"
            
            qr = qrcode.QRCode(
                version=1,
                error_correction=qrcode.constants.ERROR_CORRECT_L,
                box_size=10,
                border=2
            )
            
            qr.add_data(url)
            qr.make(fit=True)
            qr_img = qr.make_image(fill_color="black", back_color="white").convert('RGB')
            
            font = ImageFont.truetype(f"{cwd}\\fonts\\Montserrat-Light.ttf", size=20)
            
            text_bbox = font.getbbox(place.place.short_name)
            text_width = text_bbox[2] - text_bbox[0]
            text_height = text_bbox[3] - text_bbox[1]
            
            img_width = max(qr_img.width, text_width + 20)
            img_height = qr_img.height + text_height + 20
            
            img = Image.new('RGB', (img_width, img_height), 'white')

            # Вставляем QR-код в изображение
            img.paste(qr_img, ((img_width - qr_img.width) // 2, 0))

            # Рисуем текст под QR-кодом
            draw = ImageDraw.Draw(img)
            text_position = (((img_width - text_width) // 2) - 27, qr_img.height)
            draw.text(text_position, f"{"Стол" if place.name == "Стол" else "Комната"}: " 
                      + place.place.short_name, fill="black", font=font)

            # Сохраняем изображение в файл
            try:
                img.save(f"{cwd}\\qr-codes\\office {office.pk}\\{place.place.short_name}.png")
                data.append({'image': f"{cwd}\\qr-codes\\office {office.pk}\\{place.place.short_name}.png", 
                             'office': office.address, 'place': place.place.short_name})
            except Exception as ex:
                log("Произошла ошибка при создании QR-кодов!")
    
    workbook = Workbook()
    worksheet = workbook.active
    
    for row_num, d in enumerate(data, start=1):
        excel_img = ExcelImage(d['image'])
        
        worksheet.add_image(excel_img, f"A{row_num}")
        worksheet.row_dimensions[row_num].height = excel_img.height
        worksheet.column_dimensions["A"].width = 40
        worksheet.cell(row=row_num, column=2, value=d['office'])
        worksheet.cell(row=row_num, column=3, value=d['place'])
        
    
    workbook.save("QRcodes.xlsx")
    log(f"Пользователь '{user.login}' создал QR-коды!")
    return FileResponse(open("QRcodes.xlsx", 'rb'), as_attachment=True)


@head_decorator
def head_panel(request):
    """
    Returns render of a head panel.
    """
    return render(request, "head_panel.html")


@authorized
def settings(request):
    """
    Returns render of a settings page.
    """
    return render(request, 'settings.html')


@authorized
def continue_reservation(request, id):
    """
    :id: Id of a reservation\n
    Creates new reservation object on base of another reservation.
    Then saves it to database.
    """
    success = request.GET.get('success', '0')
    if success == "1":
        if Reservation.objects.filter(pk=id).exists() == False:
            return redirect('/user-reservations/')
        
        reservation = Reservation.objects.get(pk=id)
        _time = Time.objects.filter(time = "Вечернее").first()
        create_reservation(reservation.user, _time, reservation.status, reservation.type, reservation.date)
        return redirect('/success/')
    else:
        if Reservation.objects.filter(pk=id).exists() == False:
            return redirect('/user-reservations/')
        
        reservation = Reservation.objects.get(pk=id)
        
        if reservation.check_continue() is False:
            return redirect('/user-reservations/')
        
        time = "Вечернее"
        
        args = {"rs": reservation, "time": time}
        return render(request, "continue.html", args)


@authorized
def cancel(request, id):
    """
    :id: Id of a reservation to cancel \n
    Changes status of a reservation to "Canceled"
    """
    success = request.GET.get('success', '0')
    if success == "1":
        
        if Reservation.objects.filter(pk=id).exists() == False:
            return redirect('/user-reservations/')
        
        reservation = Reservation.objects.get(pk=id)
        status = Status.objects.filter(status = "Отменено").first()
        reservation.status = status
        reservation.cancel_reason = "Отменено пользователем!"
        reservation.save()
        return redirect('/success/')
        
    else:
        if Reservation.objects.filter(pk=id).exists() == False:
            return redirect('/user-reservations/')
        
        reservation = Reservation.objects.get(pk=id)
        args = {"rs": reservation}
        return render(request, "cancel.html", args)

# ///////////////////////////////////////////////////////////////

@authorized
def user_reservations(request):
    
    user_id = request.session.get('id', -1)
    if user_id == -1:
        return redirect('/login/')
    user = User.objects.filter(id=user_id).first()
    if Status.objects.filter(status = "Забронировано").exists() is False:
            status = Status()
            status.status = "Забронировано"
            status.save()
            log(f"Статус '{status.status}' создан!")
            
    if Status.objects.filter(status = "Завершено").exists() is False:
            status = Status()
            status.status = "Завершено"
            status.save()
            log(f"Статус '{status.status}' создан!")
        
    if Status.objects.filter(status = "Отменено").exists() is False:
        status = Status()
        status.status = "Отменено"
        status.save()
        log(f"Статус '{status.status}' создан!")
            
    status = Status.objects.filter(status = "Забронировано").first()
    closed = Status.objects.filter(status = "Завершено").first()
    canceled = Status.objects.filter(status = "Отменено").first()
    
    active_reservations = list(Reservation.objects.exclude(is_archivated=True) 
                               .filter(user=user, status=status).all())
    history = list(Reservation.objects.exclude(is_archivated=True).filter(user=user, status__in=[closed, canceled]).all())
    active = []
    
    for rs in active_reservations:
        active.append({"reservation": rs, "continue": rs.check_continue()})
    
    args = {'active': active, 'history': history}
    
    return render(request, 'user_reservations.html', args)

@authorized
def success(request):
    return render(request, "success.html")

@authorized
def book(request, id):
    office_id = request.GET.get('office', 'none')
    date = request.GET.get('date', 'none')
    time = request.GET.get('time', 'none')
    success = request.GET.get('success', '0')
    
    if office_id == 'none' or date == 'none' or time == 'none':
        return redirect('/reservations/')
        
    if success == "1":
        
        office = Office.objects.get(pk=int(office_id))
        date = datetime.datetime.strptime(date, "%d.%m.%Y").date()
        place = Type.objects.get(pk=id)
        
        user_id = request.session.get('id', -1)
        if user_id == -1:
            return redirect('/login/')
        user = User.objects.filter(id=user_id).first()
        
        reservation = Reservation()
        reservation.user = user
        if Time.objects.filter(time = time).exists() is False:
            time__ = Time()
            time__.time = time
            time__.save()

        time_ = Time.objects.filter(time = time).first()
        reservation.time = time_
        
        if Status.objects.filter(status = "Забронировано").exists() is False:
            status = Status()
            status.status = "Забронировано"
            status.save()
            log(f"Статус '{status.status}' создан!")
            
        status = Status.objects.filter(status = "Забронировано").first()
        reservation.status = status
        reservation.type = place
        reservation.date = date
        reservation.is_archivated = False
        reservation.save()
        
        return redirect('/success/')
    
    office = Office.objects.get(pk=int(office_id))
    date = datetime.datetime.strptime(date, "%d.%m.%Y")
    
    place = Type.objects.get(pk=id)
    
    args = {'time': time, 'office_str': office, 'date_str': date.date, 'place': place, 'office': office.pk, 'date': date.strftime("%d.%m.%Y")}
    
    return render(request, 'finish_booking.html', args)


@authorized
def reservations(request):
    
    office = Office.objects.first()
    date = datetime.date.today()
    time = "Утреннее" if datetime.datetime.now().time() < datetime.time(14, 0) else "Вечернее"
    
    rooms = Type.objects.filter(name="Комната", place__office=office)
    tables = Type.objects.filter(name="Стол", place__office=office)
    
    if request.method == "POST":
        form = ReservationFilterForm(request.POST)
        if form.is_valid():
            office = form.cleaned_data['office']
            date = form.cleaned_data['date']
            time = form.cleaned_data['time']
    
    else:
        form = ReservationFilterForm(initial={"office": office, "date": date, "time": time})
    
    status_cancelled = Status.objects.filter(status = "Отменено").first()
    status_closed = Status.objects.filter(status="Завершено").first()
    
    reservations = Reservation.objects.exclude(is_archivated=True).filter(date=date, time__time=time).exclude(status__in=[status_closed, status_cancelled]).all()
    
    rooms_available = []
    tables_available = []
    
    rooms_unavailable = []
    tables_unavailable = []
    
    for room in rooms:
        available = True
        
        for rs in reservations:
            if room == rs.type:
                rooms_unavailable.append(room)
                available = False
                break
        
        if available:
            rooms_available.append(room)
    
    
    for table in tables:
        available = True
        
        for rs in reservations:
            if table == rs.type:
                tables_unavailable.append(table)
                available = False
                break
        
        if available:
            tables_available.append(table)
        
    args = {'form': form, "rooms_available": rooms_available, "table_available": tables_available, "rooms_unavailable": rooms_unavailable, "table_unavailable": tables_unavailable, 'office': office.pk, 'date': date.strftime("%d.%m.%Y"), 'time': time}
    
    return render(request, 'reservations.html', args)


@authorized
def menu(request):
    role = request.session.get('role', 'undefined')
    args = {'role': role}
    return render(request, 'menu.html', args)


@admin_decorator
def set_archive_time(request):
    if request.method != "POST":
        return redirect('/admin/')
    
    form = TimeToArchivateForm(request.POST)
    if form.is_valid():
        setting = Setting.objects.first()
        setting.months_to_archive = form.cleaned_data['months_to_archive']
        setting.save()
        return HttpResponseRedirect('/admin/?entity=settings')
    

@admin_decorator
def add_entity(request):
    
    if request.method == "POST":
        entity = request.POST.get('entity', 'none')
    else:
        entity = request.GET.get('entity', 'none')
    
    if entity == 'none':
        return redirect('/admin/')
    
    if entity == "users":
        form = UserFormAdd() if request.method != "POST" else UserFormAdd(request.POST)
    elif entity == "offices":
        form = OfficeForm() if request.method != "POST" else OfficeForm(request.POST)
    elif entity == "places":
        form = PlaceForm() if request.method != "POST" else PlaceForm(request.POST)
    elif entity == "reservations":
        form = ReservationForm() if request.method != "POST" else ReservationForm(request.POST)
    
    if request.method == "POST":
        if form.is_valid():
            form.save(form.cleaned_data)
            return redirect(f'/admin/?entity={entity}')
    
    args = {'entity': entity, 'form': form}
    
    return render(request, "admin_add_entity.html", args)

@admin_decorator
def create_dump(request):
    id = request.session.get('id', -1)
    if id == -1:
        return redirect('/login/')
    
    setting = Setting.objects.first()
    two_months_ago = datetime.datetime.now() - relativedelta(months=setting.months_to_archive)
    
    reservations = Reservation.objects.filter(creation_time__lt=two_months_ago).exclude(is_archivated=True).all()
    
    for rs in reservations:
        rs.is_archivated = True
        rs.save()
        
    archived = Reservation.objects.filter(is_archivated=True).all()
    
    from django.core import serializers
    json = serializers.serialize('json', archived)
    
    with open('archive.json', 'w', encoding="UTF-8") as f:
        f.write(json)
    
    user = User.objects.filter(id = id).first()
    log(f"Пользователь '{user.login}' создал архив старых значений!")
    log(f"Пользователь '{user.login}' создал дамп БД!")
    buf = StringIO()
    call_command('dumpdata', stdout=buf)
    buf.seek(0)
    with open('dump.json', 'w', encoding="UTF-8") as f:
        f.write(buf.read())
    return FileResponse(open('dump.json', 'rb'), as_attachment=True)


@head_decorator
def download_report(request):
    reservations = list(Reservation.objects.filter(is_archivated=False).all())
    id = request.session.get('id', -1)
    if id == -1:
        return redirect('/login/')
    user = User.objects.filter(id = id).first()
    
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    document = Document()
    h = document.add_heading("ОТЧЕТ", 1)
    p1 = document.add_paragraph("О всех бронированиях в системе бронирований Bookingz")
    p2 = document.add_paragraph('ООО "vsyakoe"')
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table = document.add_table(rows=1, cols=9)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Пользователь'
    hdr_cells[2].text = 'Место'
    hdr_cells[3].text = 'Тип места'
    hdr_cells[4].text = 'Дата бронирования'
    hdr_cells[5].text = 'Время бронирования'
    hdr_cells[6].text = 'Статус'
    hdr_cells[7].text = 'Причина отмены'
    hdr_cells[8].text = 'Время создания'
    for item in reservations:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item.pk)
        row_cells[1].text = f"{item.user.name} {item.user.lastname}"
        row_cells[2].text = str(item.type.place.short_name)
        row_cells[3].text = str(item.type.name)
        row_cells[4].text = item.date.strftime("%d.%m.%Y")
        row_cells[5].text = str(item.time.time)
        row_cells[6].text = str(item.status.status)
        row_cells[7].text = str(item.cancel_reason)
        row_cells[8].text = str(item.creation_time)
       
    table.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
       
    p3 = document.add_paragraph(f"Руководитель: {user.name} {user.lastname}") 
    p4 = document.add_paragraph("_______________ Подпись")
    
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    document.save('report.docx')
    return FileResponse(open('report.docx', 'rb'), as_attachment=True)

    

@admin_decorator
def ban_user(request):
    id = int(request.GET.get('id', "-1"))
    if id == -1:
        return redirect('/admin/')
    user = User.objects.get(id=id)
    user.is_blocked = True
    log(f"Пользователь с id: '{user.pk}' заблокирован!")
    user.save()
    return redirect('/admin/')

@admin_decorator
def unban_user(request):
    id = int(request.GET.get('id', "-1"))
    if id == -1:
        return redirect('/admin/')
    user = User.objects.get(id=id)
    user.is_blocked = False
    user.save()
    log(f"Пользователь с id: '{user.pk}' разблокирован!")
    return redirect('/admin/')

@admin_decorator
def cancel_reservation(request):
    id = int(request.GET.get('id', "-1"))
    if id == -1:
        return redirect('/admin/')
    rs = Reservation.objects.get(id=id)
    status = Status.objects.filter(status = "Отменено").first()
    rs.status = status
    rs.cancel_reason = "Отменено администратором"
    rs.save()
    log(f"Бронирование с id: '{rs.pk}' отменено администратором!")
    return redirect(f'/admin/?entity=reservations')
    
@admin_decorator
def edit_entity(request):    
    
    if request.method == "POST":
        entity = request.POST.get('entity', 'none')
        id = int(request.POST.get('id', '-1'))
    else:
        entity = request.GET.get('entity', 'none')
        id = int(request.GET.get('id', '-1'))
    
    if entity == 'none' or id == -1:
        return redirect(f'/admin/?entity={entity}')
    
    if entity == "users":
        entity_inst = User.objects.get(id=id)
        form = UserForm(instance=entity_inst) if request.method != "POST" else UserForm(request.POST)
    elif entity == "offices":
        entity_inst = Office.objects.get(id=id)
        form = OfficeForm(instance=entity_inst) if request.method != "POST" else OfficeForm(request.POST)
    elif entity == "places":
        entity_inst = Place.objects.get(id=id)
        form = PlaceForm() if request.method != "POST" else PlaceForm(request.POST)
    elif entity == "reservations":
        entity_inst = Reservation.objects.get(id=id)
        form = ReservationForm(instance=entity_inst) if request.method != "POST" else ReservationForm(request.POST)
    
    if request.method == "POST":
        if form.is_valid():
            form.update(form.cleaned_data, id)
            log(f"Объект {entity} с id: '{id}' обновлен!")
            return redirect(f'/admin/?entity={entity}')
    
    args = {'entity': entity, 'form': form, 'id': id}
        
    return render(request, "admin_edit_entity.html", args)
    
@admin_decorator
def delete_entity(request):
    entity = request.GET.get('entity', 'none')
    id = request.GET.get('id', '-1')
    
    if entity == 'none' or id == '-1':
        return redirect(f'/admin/?entity={entity}')
    
    if entity == "users":
        entity_inst = User.objects.get(id=id)
        log(f"Пользователь с id: '{id}' удален!")
        entity_inst.delete()
        return redirect(f'/admin/?entity={entity}')
        
    elif entity == "offices":
        entity_inst = Office.objects.get(id=id)
        log(f"Офис с id: '{id}' удален!")
        entity_inst.delete()
        return redirect(f'/admin/?entity={entity}')
        
    elif entity == "places":
        entity_inst = Place.objects.get(id=id)
        log(f"Место с id: '{id}' удалено!")
        entity_inst.delete()
        return redirect(f'/admin/?entity={entity}')
        
        
    elif entity == "reservations":
        entity_inst = Reservation.objects.get(id=id)
        log(f"Бронирование с id: '{id}' удалено!")
        entity_inst.delete()
        return redirect(f'/admin/?entity={entity}')

@admin_decorator
@transaction.atomic
def upload(request):
    if request.method != "POST":
        return redirect('/admin/')
    
    form = UploadFileForm(request.POST, request.FILES)
    from django.core.exceptions import ValidationError
    if form.is_valid():
        f = request.FILES.get("file", None)
        if not '.json' in f.name:
            raise ValidationError('File type not supported.')
        if f is not None:
            with open("db-dump.json", "wb+") as destination:
                for chunk in f.chunks():
                    destination.write(chunk)
            from django.core.management.commands import flush
            
            Role.objects.all().delete()
            User.objects.all().delete()
            Time.objects.all().delete()
            Status.objects.all().delete()
            Office.objects.all().delete()
            Place.objects.all().delete()
            Type.objects.all().delete()
            Reservation.objects.all().delete()
            log("База данных сброшена!")
            
            call_command('loaddata', "db-dump.json")
            call_command('migrate')
            log("База данных загружена из дампа!")
            
            return HttpResponseRedirect("/admin/")
        else:
            return redirect('/admin/')
    else:
        return redirect('/admin/?entity=settings')
    

@admin_decorator
def admin_page(request):        
    entity = request.GET.get('entity', 'users')
    count = 0
    data = []
    
    if entity == "users":
        id = request.session.get('id', -1)
        if id == -1:
            return redirect('/login/')
        user = User.objects.filter(id = id).first()
        data = list(User.objects.exclude(id=user.pk))
    elif entity == "offices":
        data = list(Office.objects.all().order_by('pk'))
    elif entity == "places":
        data = list(Type.objects.all().order_by('pk'))
    elif entity == "reservations":
        with connection.cursor() as cursor:
            cursor.execute("SELECT get_reservation_counter();")
            count = cursor.fetchone()[0]
            
        data = list(Reservation.objects.exclude(is_archivated=True).all().order_by('pk'))
    elif entity == "logs":
        data = list(Log.objects.all().order_by('pk'))
    elif entity == 'graph':
        data_points = []
        graphic = request.GET.get('graphic', 'days')
        chart_title = ""
        data_type = "column"
        
        if graphic == "days":
            chart_title = "График бронирований по дням"
            today = datetime.date.today()
            weekday = today.weekday()
            start_delta = datetime.timedelta(days=weekday, weeks=1)
            start_of_week = today - start_delta
            for day in range(7):
                date = start_of_week + datetime.timedelta(days=day)
                exclude_status = Status.objects.filter(status = "Отменено").first()
                reservations = Reservation.objects.filter(date=date).exclude(status=exclude_status).all()
                days_name = ['Понедельник', 'Вторник', 'Среда', 'Четверг', 'Пятница', 'Суббота', 'Воскресенье']
                data_points.append({"label": days_name[day], "y": len(reservations)})
        
        elif graphic == "tables":
            chart_title = "График 10 самых популярных столов по бронированиям"
            exclude_status = Status.objects.filter(status = "Отменено").first()
            reservations = Reservation.objects.exclude(status=exclude_status).exclude(is_archivated=True).all()
            from collections import Counter, OrderedDict
            tables = []
            for rs in reservations:
                if rs.type.name == "Стол":
                    tables.append(rs.type.pk)
            
            cnt = dict(Counter(tables))
            cnt = OrderedDict(sorted(cnt.items(), key=lambda kv: kv[1], reverse=True)[:10])
            
            for key in cnt:
                type_ = Type.objects.get(pk=key)
                data_points.append({'label': type_.place.short_name, 'y': cnt[key]})
        
        else:
            chart_title = "График пользования комнатами"
            data_type = "area"
            exclude_status = Status.objects.filter(status = "Отменено").first()
            reservations = Reservation.objects.exclude(status=exclude_status).exclude(is_archivated=True).all()
            from collections import Counter, OrderedDict
            rooms = []
            for rs in reservations:
                if rs.type.name == "Комната":
                    rooms.append(rs.date.strftime("%d.%m.%Y"))
                    
            cnt = dict(Counter(rooms))

            for key in cnt:
                data_points.append({'label': key, 'y': cnt[key]})
        
        data = data_points
        
    elif entity == 'archive':
        
        reservations = []
        
        if os.path.exists('archive.json'):
            with open('archive.json', 'r', encoding="UTF-8") as file:
                json_ = json.load(file)
            
            for item in json_:
                rs = dict()
                __user = User.objects.get(pk=item['fields']['user'])
                time = Time.objects.get(pk=item['fields']['time'])
                status = Status.objects.get(pk=item['fields']['status'])
                __type = Type.objects.get(pk=item['fields']['type'])
                rs.update({"id": item['pk'] })
                rs.update(item['fields'])
                rs.update({"user" : __user})
                rs.update({"time": time})
                rs.update({"status": status})
                rs.update({"type": __type})            
                reservations.append(rs)
            
            data = reservations
    
    form = UploadFileForm()
    
    if Setting.objects.first() is None:
        setting = Setting()
        setting.months_to_archive = 2
        setting.save()
    
    setting = Setting.objects.first()
    initial_dict = {
        "months_to_archive": setting.months_to_archive
    }
    
    form_months = TimeToArchivateForm(initial = initial_dict)
    args = {'entity': entity, 'data': data, 'form': form, 'form_months': form_months, 'count': count}
    if entity == "graph":
        args.update({"graphic": graphic})
        args.update({'chart_title': chart_title})
        args.update({'data_type': data_type})
    return render(request, 'admin_page.html', args)


def homepage(request):
    return render(request, "homepage.html")


def reg_page(request):
    if Role.objects.filter(name = "rseU").exists() is False:
        role = Role()
        role.name = "User"
        role.crypt()
        role.save()
        log("Роль 'User' создана!")
    
    role = Role.objects.filter(name = "rseU").get()
    
    existing_role = request.session.get('role', 'undefined')
    if existing_role != "undefined":
        return redirect('/menu/')
    
    if request.method == "POST":
        form = RegForm(request.POST)
        if form.is_valid():
            data = form.cleaned_data
            user = User()
            user.login = data['login']
            user.email = data['email']
            hasher = hashlib.sha256(usedforsecurity=True)
            hasher.update(str(data['password']).encode())
            h_password = hasher.hexdigest()
            user.password = h_password
            user.name = data['name']
            user.lastname = data['lastname']
            user.role = role
            user.is_blocked = False
            user.save()
            request.session['id'] = user.pk
            request.session['role'] = user.role.decrypt()
            return redirect('/menu/')
    
    else:
        form = RegForm()
    
    return render(request, 'reg.html', {"form": form})
    


def login_page(request):
    if Role.objects.filter(name = "ndmiA").exists() is False:
        role = Role()
        role.name = "Admin"
        role.crypt()
        role.save()
        log("Роль 'Admin' создана!")
    
    role = Role.objects.filter(name = "ndmiA").get()
    
    if User.objects.filter(role = role).exists() is False:
        admin = User()
        admin.name = "admin"
        admin.lastname = "admin"
        admin.email = "admin@email.com"
        admin.login = "admin"
        hasher = hashlib.sha256(usedforsecurity=True)
        hasher.update(str("admin").encode())
        h_password = hasher.hexdigest()
        admin.password = h_password
        role = role
        admin.role = role
        admin.is_blocked = False
        admin.save()
        log(f"Пользователь с id: '{admin.pk}' и ролью: '{admin.role.decrypt()}' создан!")
        
    existing_role = request.session.get('role', 'undefined')
    if existing_role != "undefined":
        return redirect('/menu/')
    
    if request.method == "POST":
        form = LoginForm(request.POST)
        if form.is_valid():
            data = form.cleaned_data
            hasher = hashlib.sha256(usedforsecurity=True)
            hasher.update(str(data['password']).encode())
            h_password = hasher.hexdigest()
            if User.objects.filter(login = data['login']).filter(password = h_password).filter(is_blocked = False).exists():
                user = User.objects.filter(login = data['login']).filter(password = h_password).filter(is_blocked = False).get()
                request.session['id'] = user.pk
                request.session['role'] = user.role.decrypt()
                if user.role.name == "Admin":
                    log(f"Пользователь с id: '{user.pk}' с ролью: '{user.role.decrypt()}' вошел в систему!")
                    return redirect('/admin/')
                else:
                    log(f"Пользователь с id: '{user.pk}' с ролью: '{user.role.decrypt()}' вошел в систему!")
                    return redirect('/menu/')
    
    else:
        form = LoginForm()
    
    return render(request, 'login.html', {"form": form})


@authorized
def profile(request):
    id = request.session.get('id', -1)
    if id == -1:
        return redirect('/login/')
    user = User.objects.filter(id = id).first()
    args = {'user': user}
    return render(request, 'profile.html', args)

@authorized
def logout(request):
    id = request.session.get('id', -1)
    if id == -1:
        return redirect('/login/')
    user = User.objects.filter(id = id).first()
    request.session.flush()
    log(f"{user.login} вышел!")
    return redirect('/login/')

@authorized
def update_user(request):
    id = request.session.get('id', -1)
    if id == -1:
        return redirect('/login/')
    
    user = User.objects.filter(id = id).first()
    
    if request.method == "POST":
        form = UpdateProfile(request.POST)
        if form.is_valid():
            form.update(form.cleaned_data, id)
            return redirect('/profile/')
    
    else:
        form = UpdateProfile(instance=user)
        
    args = {'id': id, 'form': form}
    
    return render(request, 'profile_change.html', args)


@authorized
def delete_account(request):
    id = request.session.get('id', -1)
    if id == -1:
        return redirect('/login/')
    user = User.objects.filter(id = id).first()
    log(f"{user.login} удален!")
    user.delete()
    request.session.flush()
    return redirect('/login/')